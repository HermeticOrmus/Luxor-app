import fitz  # PyMuPDF for PDFs
from docx import Document as DocxReader  # python-docx for .docx
from llama_index.core.schema import Document

def load_pdf_with_metadata(file_path):
    """
    Loads a PDF file page by page, chunks it into 5-line segments,
    and attaches metadata: file name, page number, and starting line.
    """
    doc = fitz.open(file_path)
    documents = []
    file_name = file_path.split("/")[-1]

    for page_number, page in enumerate(doc):
        text = page.get_text("text")
        lines = text.splitlines()

        for i in range(0, len(lines), 5):  # Chunk every 5 lines
            chunk = "\n".join(lines[i:i+5])
            if not chunk.strip():
                continue

            metadata = {
                "file_name": file_name,
                "page_number": page_number + 1,
                "line_start": i + 1
            }

            documents.append(Document(text=chunk, metadata=metadata))
    
    return documents


def load_docx_with_metadata(file_path):
    """
    Loads a Word .docx file, chunks it into 3-paragraph segments,
    and attaches metadata: file name and paragraph start index.
    (No page support for .docx format)
    """
    reader = DocxReader(file_path)
    paragraphs = [p.text.strip() for p in reader.paragraphs if p.text.strip()]
    
    documents = []
    file_name = file_path.split("/")[-1]

    for i in range(0, len(paragraphs), 3):
        chunk = "\n".join(paragraphs[i:i+3])
        if not chunk.strip():
            continue

        metadata = {
            "file_name": file_name,
            "page_number": "N/A",
            "line_start": i + 1
        }

        documents.append(Document(text=chunk, metadata=metadata))
    
    return documents
